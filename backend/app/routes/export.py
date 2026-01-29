"""
Export Routes - DOCX and PDF generation
"""

from fastapi import APIRouter, HTTPException
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import List
import io
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER

router = APIRouter()

class ExportRequest(BaseModel):
    repo_name: str
    repo_owner: str
    explanation: dict
    resume_bullets: List[dict]
    viva_questions: List[dict]
    interview_qa: List[dict]

@router.post("/export-docx")
async def export_docx(data: ExportRequest):
    """Export analysis as DOCX file"""
    
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading(f'{data.repo_owner}/{data.repo_name}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading('GitHub Repository Analysis', level=2)
        doc.add_paragraph()
        
        # Project Explanation
        doc.add_heading('📖 Project Overview', level=1)
        doc.add_paragraph(data.explanation['overview'])
        doc.add_paragraph()
        
        doc.add_heading('Key Features', level=2)
        for feature in data.explanation['key_features']:
            doc.add_paragraph(feature, style='List Bullet')
        doc.add_paragraph()
        
        doc.add_heading('Tech Stack', level=2)
        doc.add_paragraph(', '.join(data.explanation['tech_stack']))
        doc.add_paragraph()
        
        doc.add_heading('Architecture', level=2)
        doc.add_paragraph(data.explanation['architecture'])
        doc.add_paragraph()
        
        doc.add_heading('Challenges Solved', level=2)
        for challenge in data.explanation['challenges_solved']:
            doc.add_paragraph(challenge, style='List Bullet')
        doc.add_paragraph()
        
        doc.add_heading('Impact', level=2)
        doc.add_paragraph(data.explanation['impact'])
        
        # Page break
        doc.add_page_break()
        
        # Resume Bullets
        doc.add_heading('📄 Resume Bullet Points', level=1)
        for bullet in data.resume_bullets:
            doc.add_paragraph(bullet['point'], style='List Bullet')
        
        # Page break
        doc.add_page_break()
        
        # Viva Questions
        doc.add_heading('🎓 Viva Questions', level=1)
        for i, viva in enumerate(data.viva_questions, 1):
            doc.add_heading(f'Q{i} [{viva["difficulty"].upper()}]', level=2)
            doc.add_paragraph(f'Question: {viva["question"]}')
            doc.add_paragraph(f'Answer: {viva["answer"]}')
            doc.add_paragraph()
        
        # Page break
        doc.add_page_break()
        
        # Interview Q&A
        doc.add_heading('💼 Interview Questions & Answers', level=1)
        for i, qa in enumerate(data.interview_qa, 1):
            doc.add_heading(f'Q{i} [{qa["category"].upper()}]', level=2)
            doc.add_paragraph(f'Question: {qa["question"]}')
            doc.add_paragraph(f'Answer: {qa["answer"]}')
            doc.add_paragraph()
        
        # Save to BytesIO
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        return StreamingResponse(
            file_stream,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={data.repo_owner}-{data.repo_name}-analysis.docx"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DOCX export failed: {str(e)}")


@router.post("/export-pdf")
async def export_pdf(data: ExportRequest):
    """Export analysis as PDF file"""
    
    try:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch)
        
        styles = getSampleStyleSheet()
        story = []
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=RGBColor(0, 0, 139),
            alignment=TA_CENTER,
            spaceAfter=20
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=RGBColor(0, 0, 139),
            spaceAfter=12
        )
        
        # Title
        story.append(Paragraph(f'{data.repo_owner}/{data.repo_name}', title_style))
        story.append(Paragraph('GitHub Repository Analysis', styles['Heading2']))
        story.append(Spacer(1, 0.3*inch))
        
        # Project Overview
        story.append(Paragraph('📖 Project Overview', heading_style))
        story.append(Paragraph(data.explanation['overview'], styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph('Key Features', styles['Heading3']))
        for feature in data.explanation['key_features']:
            story.append(Paragraph(f'• {feature}', styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph('Tech Stack', styles['Heading3']))
        story.append(Paragraph(', '.join(data.explanation['tech_stack']), styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph('Architecture', styles['Heading3']))
        story.append(Paragraph(data.explanation['architecture'], styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph('Challenges Solved', styles['Heading3']))
        for challenge in data.explanation['challenges_solved']:
            story.append(Paragraph(f'• {challenge}', styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        story.append(Paragraph('Impact', styles['Heading3']))
        story.append(Paragraph(data.explanation['impact'], styles['Normal']))
        
        story.append(PageBreak())
        
        # Resume Bullets
        story.append(Paragraph('📄 Resume Bullet Points', heading_style))
        for bullet in data.resume_bullets:
            story.append(Paragraph(f'• {bullet["point"]}', styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        
        story.append(PageBreak())
        
        # Viva Questions
        story.append(Paragraph('🎓 Viva Questions', heading_style))
        for i, viva in enumerate(data.viva_questions, 1):
            story.append(Paragraph(f'Q{i} [{viva["difficulty"].upper()}]', styles['Heading3']))
            story.append(Paragraph(f'<b>Question:</b> {viva["question"]}', styles['Normal']))
            story.append(Paragraph(f'<b>Answer:</b> {viva["answer"]}', styles['Normal']))
            story.append(Spacer(1, 0.15*inch))
        
        story.append(PageBreak())
        
        # Interview Q&A
        story.append(Paragraph('💼 Interview Questions & Answers', heading_style))
        for i, qa in enumerate(data.interview_qa, 1):
            story.append(Paragraph(f'Q{i} [{qa["category"].upper()}]', styles['Heading3']))
            story.append(Paragraph(f'<b>Question:</b> {qa["question"]}', styles['Normal']))
            story.append(Paragraph(f'<b>Answer:</b> {qa["answer"]}', styles['Normal']))
            story.append(Spacer(1, 0.15*inch))
        
        doc.build(story)
        buffer.seek(0)
        
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={data.repo_owner}-{data.repo_name}-analysis.pdf"}
        )
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"PDF export failed: {str(e)}")